#!/usr/bin/env python3
import os
import sys
import json
import asyncio
import base64
import argparse
from typing import List, Dict, Optional
from pydantic import BaseModel, Field
from typing_extensions import Literal

# Ensure packages can be imported
try:
    import fitz  # PyMuPDF
except ImportError:
    print("Error: PyMuPDF is not installed. Run 'pip install pymupdf'")
    sys.exit(1)

try:
    import docx
except ImportError:
    print("Error: python-docx is not installed. Run 'pip install python-docx'")
    sys.exit(1)

try:
    import openpyxl
except ImportError:
    print("Error: openpyxl is not installed. Run 'pip install openpyxl'")
    sys.exit(1)

try:
    import edge_tts
except ImportError:
    print("Error: edge-tts is not installed. Run 'pip install edge-tts'")
    sys.exit(1)

try:
    from google import genai
    from google.genai import types
except ImportError:
    print("Error: google-genai is not installed. Run 'pip install google-genai'")
    sys.exit(1)



# Pydantic Schemas for Structured output from Gemini API
class QuestionSchema(BaseModel):
    id: int = Field(description="Unique question index starting from 1.")
    type: Literal["multiple-choice", "yes-no", "drag-and-drop"] = Field(
        description="The style of interaction: multiple-choice, yes-no, or drag-and-drop matching."
    )
    question: str = Field(description="The question prompt or matcher instructions.")
    options: Optional[List[str]] = Field(
        default=None,
        description="List of options to display. For multiple-choice, supply 3-5 options. For yes-no, supply ['Yes', 'No']. Leave empty/null for drag-and-drop matching.",
    )
    answer: Optional[str] = Field(
        default=None,
        description="The exact string answer from the options list. Required for multiple-choice and yes-no. Leave empty/null for drag-and-drop.",
    )
    pairs: Optional[Dict[str, str]] = Field(
        default=None,
        description="Dictionary of matching items (e.g. {'France': 'Paris', 'Germany': 'Berlin'}). Required ONLY if type is 'drag-and-drop'. Leave empty/null otherwise.",
    )
    explanation: str = Field(
        description="Detailed context explaining why the answer is correct or why the matched items pair together."
    )


class QuizSchema(BaseModel):
    title: str = Field(description="Catchy title reflecting the content of the document.")
    description: str = Field(description="A brief description of what this quiz covers and its difficulty level.")
    questions: List[QuestionSchema] = Field(description="List of questions compiled from the document.")


# Parsing functions
def parse_pdf(file_path: str) -> str:
    """Extracts raw text content from PDF pages."""
    text_content = []
    with fitz.open(file_path) as doc:
        for page_num, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                text_content.append(page_text)
    
    if not text_content:
        raise ValueError("The PDF document seems to be empty or contains no extractable text.")
    return "\n\n--- Page Break ---\n\n".join(text_content)


def parse_docx(file_path: str) -> str:
    """Extracts text paragraphs and table rows from Word Document."""
    doc = docx.Document(file_path)
    text_content = []
    
    # Extract Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            text_content.append(p.text.strip())
            
    # Extract Tables
    for idx, table in enumerate(doc.tables):
        text_content.append(f"\n[Table {idx+1}]")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                # remove duplicates created by merged cells
                unique_row = []
                for val in row_text:
                    if not unique_row or unique_row[-1] != val:
                        unique_row.append(val)
                text_content.append(" | ".join(unique_row))
                
    if not text_content:
        raise ValueError("The DOCX document seems to be empty.")
    return "\n".join(text_content)


def parse_xlsx(file_path: str) -> str:
    """Extracts row data and sheet information from Excel file."""
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    text_content = []
    
    for sheet_name in wb.sheetnames:
        text_content.append(f"\n--- Sheet: {sheet_name} ---")
        ws = wb[sheet_name]
        for row in ws.iter_rows(values_only=True):
            # Format row elements to string, filter empty rows
            row_vals = [str(val).strip() for val in row if val is not None]
            if row_vals:
                text_content.append(" | ".join(row_vals))
                
    if not text_content:
        raise ValueError("The Excel workbook is empty.")
    return "\n".join(text_content)


def extract_document_text(file_path: str) -> str:
    """Detects file extension and routes to appropriate parser."""
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == '.pdf':
        return parse_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return parse_docx(file_path)
    elif ext in ['.xlsx', '.xls']:
        return parse_xlsx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and XLSX are supported.")


# LLM integration
def generate_quiz_json(document_text: str, api_key: str, num_questions: int) -> str:
    """Sends document contents to Gemini to generate quiz JSON."""
    client = genai.Client(api_key=api_key)
    
    prompt = f"""
Analyze the following document and generate a quiz to test the reader's understanding.
Generate exactly {num_questions} questions.

You must design a mix of the following question types:
1. "multiple-choice": Question with 3 to 5 options. Ensure only ONE option is correct.
2. "yes-no": A simple binary true/false style question. Use options: ["Yes", "No"].
3. "drag-and-drop": A matching exercise. Define a mapping in `pairs` where the keys are items to drag and values are their corresponding targets. Ensure there are 3 to 5 pairs for this question. Do not supply options or answer attributes for drag-and-drop.

For every single question, provide a solid explanation of the correct answer based on facts directly found in the document.

DOCUMENT CONTENT:
{document_text}
"""
    
    # Use gemini-2.5-flash for speed, cost, and high schema accuracy
    response = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=prompt,
        config=types.GenerateContentConfig(
            response_mime_type="application/json",
            response_schema=QuizSchema,
            temperature=0.2
        )
    )
    
    return response.text


async def text_to_base64_audio(text: str, voice: str) -> str:
    """Generate TTS base64 audio string using edge-tts."""
    if not text or not text.strip():
        return ""
    try:
        # Strip simple HTML formatting to avoid reading tags
        clean_text = text.replace("<strong>", "").replace("</strong>", "").replace("<br>", " ").replace("<br><br>", " ")
        communicate = edge_tts.Communicate(clean_text, voice)
        audio_data = bytearray()
        async for chunk in communicate.stream():
            if chunk["type"] == "audio":
                audio_data.extend(chunk["data"])
        if not audio_data:
            return ""
        b64_audio = base64.b64encode(audio_data).decode('utf-8')
        return f"data:audio/mp3;base64,{b64_audio}"
    except Exception as e:
        print(f"[-] Warning: Failed to generate audio for text '{text[:20]}...': {e}")
        return ""


async def generate_quiz_audios_async(quiz_data: dict, voice: str) -> dict:
    """Generates base64 audio for all texts in the quiz concurrently."""
    tasks = []
    
    # 1. Title and description
    tasks.append(("title", text_to_base64_audio(quiz_data.get("title", ""), voice)))
    tasks.append(("description", text_to_base64_audio(quiz_data.get("description", ""), voice)))
    
    # 2. Questions
    for q_idx, q in enumerate(quiz_data.get("questions", [])):
        tasks.append((f"q_{q_idx}_text", text_to_base64_audio(q.get("question", ""), voice)))
        tasks.append((f"q_{q_idx}_explanation", text_to_base64_audio(q.get("explanation", ""), voice)))
        
        # Options
        for opt_idx, opt in enumerate(q.get("options") or []):
            tasks.append((f"q_{q_idx}_opt_{opt_idx}", text_to_base64_audio(opt, voice)))
            
        # Pairs (drag-and-drop)
        pairs = q.get("pairs") or {}
        for pair_key, pair_val in pairs.items():
            tasks.append((f"q_{q_idx}_pair_k_{pair_key}", text_to_base64_audio(pair_key, voice)))
            tasks.append((f"q_{q_idx}_pair_v_{pair_val}", text_to_base64_audio(pair_val, voice)))

    # Gather results concurrently
    keys = [t[0] for t in tasks]
    coros = [t[1] for t in tasks]
    results = await asyncio.gather(*coros)
    
    audio_map = dict(zip(keys, results))
    
    # Map results back to quiz_data
    quiz_data["title_audio"] = audio_map["title"]
    quiz_data["description_audio"] = audio_map["description"]
    
    for q_idx, q in enumerate(quiz_data.get("questions", [])):
        q["question_audio"] = audio_map[f"q_{q_idx}_text"]
        q["explanation_audio"] = audio_map[f"q_{q_idx}_explanation"]
        
        if q.get("options"):
            q["options_audio"] = [
                audio_map[f"q_{q_idx}_opt_{opt_idx}"] for opt_idx in range(len(q["options"]))
            ]
            
        if q.get("pairs"):
            q["pairs_audio"] = {
                k: {
                    "key_audio": audio_map[f"q_{q_idx}_pair_k_{k}"],
                    "val_audio": audio_map[f"q_{q_idx}_pair_v_{v}"]
                } for k, v in q["pairs"].items()
            }
            
    return quiz_data


def build_html_quiz(quiz_json_str: str, template_path: str, output_path: str):
    """Injects generated Quiz JSON into the HTML template."""
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template HTML file not found at: {template_path}")
        
    with open(template_path, 'r', encoding='utf-8') as f:
        template_html = f.read()
        
    # Replace the placeholder comment with the raw JSON string
    injected_html = template_html.replace('/* {{QUIZ_DATA_PLACEHOLDER}} */', quiz_json_str)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(injected_html)


def main():
    parser = argparse.ArgumentParser(
        description="Generate an interactive HTML/CSS/JS Quiz from PDF, DOCX, or XLSX documents using Gemini."
    )
    parser.add_argument("input_file", help="Path to the input document (PDF, Word, or Excel file).")
    parser.add_argument("-o", "--output", help="Path to write the output HTML quiz file. Defaults to input name + .html")
    parser.add_argument("-k", "--api-key", help="Gemini API Key. Can also be set via GEMINI_API_KEY environment variable.")
    parser.add_argument("-n", "--num-questions", type=int, default=10, help="Number of questions to generate. Default is 10.")
    parser.add_argument("-t", "--template", help="Path to the quiz template file. If not set, loaded based on --theme.")
    parser.add_argument(
        "--theme",
        choices=["default", "kids"],
        default="kids",
        help="Select the theme template style: 'kids' (vibrant, playful for ages 10-14, default) or 'default' (dark glassmorphism)."
    )
    parser.add_argument(
        "--no-audio",
        action="store_true",
        help="Disable TTS audio generation (skips edge-tts audio rendering)."
    )
    parser.add_argument(
        "--voice",
        help="Microsoft Edge TTS voice to use. Defaults to en-US-AnaNeural (kids) or en-US-AvaNeural (default)."
    )
    
    args = parser.parse_args()
    
    # 1. Verify input file existence
    if not os.path.exists(args.input_file):
        print(f"[-] Error: Input file not found: {args.input_file}")
        sys.exit(1)
        
    # 2. Setup outputs and paths
    input_dir = os.path.dirname(os.path.abspath(args.input_file))
    base_name, _ = os.path.splitext(os.path.basename(args.input_file))
    
    output_path = args.output
    if not output_path:
        output_path = os.path.join(input_dir, f"{base_name}_quiz.html")
        
    script_dir = os.path.dirname(os.path.abspath(__file__))
    if args.template:
        template_path = args.template
    else:
        template_filename = f"quiz_template_{args.theme}.html"
        template_path = os.path.join(script_dir, template_filename)
    
    # 3. Handle API Key
    api_key = args.api_key or os.environ.get("GEMINI_API_KEY")
    if not api_key:
        print("[!] GEMINI_API_KEY environment variable or --api-key argument not provided.")
        try:
            api_key = input("[?] Enter your Gemini API Key: ").strip()
        except (KeyboardInterrupt, EOFError):
            print("\n[-] Cancelled.")
            sys.exit(1)
            
        if not api_key:
            print("[-] Error: Gemini API Key is required to run the quiz generation.")
            sys.exit(1)

    print(f"[*] Starting Quiz Generation for: {os.path.basename(args.input_file)}")
    
    # 4. Parse document
    try:
        print("[+] Extracting text from document...")
        doc_text = extract_document_text(args.input_file)
        print(f"[+] Successfully extracted {len(doc_text)} characters.")
    except Exception as e:
        print(f"[-] Error parsing document: {e}")
        sys.exit(1)
        
    # 5. Call Gemini API
    try:
        print(f"[+] Requesting Gemini to generate {args.num_questions} questions...")
        quiz_json_str = generate_quiz_json(doc_text, api_key, args.num_questions)
        
        # Verify JSON validity
        quiz_data = json.loads(quiz_json_str)
        print("[+] Quiz schema compiled successfully.")
    except Exception as e:
        print(f"[-] Error generating quiz with Gemini API: {e}")
        sys.exit(1)
        
    # 5.5 Generate Audio if requested
    if not args.no_audio:
        try:
            voice = args.voice
            if not voice:
                voice = "en-US-AnaNeural" if args.theme == "kids" else "en-US-AvaNeural"
            print(f"[+] Generating TTS audio using voice: {voice}...")
            quiz_data = asyncio.run(generate_quiz_audios_async(quiz_data, voice))
            quiz_json_str = json.dumps(quiz_data)
            print("[+] Audio generation completed.")
        except Exception as e:
            print(f"[-] Warning: Audio generation failed: {e}. Generating quiz without audio.")
            quiz_json_str = json.dumps(quiz_data)
    else:
        quiz_json_str = json.dumps(quiz_data)
        
    # 6. Ingest into template
    try:
        print("[+] Injected quiz data into HTML template...")
        build_html_quiz(quiz_json_str, template_path, output_path)
        print(f"[+] Success! Quiz file generated at: {output_path}")
        print("[*] Open the file in any browser to play.")
    except Exception as e:
        print(f"[-] Error writing final HTML file: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
